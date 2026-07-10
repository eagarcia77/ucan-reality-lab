from __future__ import annotations

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from pathlib import Path
from urllib.parse import urlparse
from PIL import Image
from pypdf import PdfReader
from docx import Document
from bs4 import BeautifulSoup
import base64
import datetime as dt
import html
import io
import json
import mimetypes
import os
import re
import shutil
import uuid
import zipfile
import xml.etree.ElementTree as ET
import requests

BASE = Path(__file__).resolve().parent
DATA = BASE / "data"
UPLOADS = DATA / "uploads"
SCORM = DATA / "scorm"
PROJECTS = DATA / "projects"
for folder in (UPLOADS, SCORM, PROJECTS):
    folder.mkdir(parents=True, exist_ok=True)

APP_NAME = os.getenv("APP_NAME", "UCAN Reality Lab v6.0")
AI_BASE_URL = os.getenv("AI_BASE_URL", "").rstrip("/")
AI_API_KEY = os.getenv("AI_API_KEY", "")
AI_MODEL = os.getenv("AI_MODEL", "")
SKETCHFAB_API_TOKEN = os.getenv("SKETCHFAB_API_TOKEN", "")

app = FastAPI(title=APP_NAME, version="6.0")
app.mount("/static", StaticFiles(directory=str(BASE / "static")), name="static")
app.mount("/data", StaticFiles(directory=str(DATA)), name="data")

ALLOWED_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".pdf", ".docx", ".txt", ".md", ".html", ".htm", ".csv"}
MAX_UPLOAD = 25 * 1024 * 1024

TOPICS = {
    "motherboard": ["placa madre", "motherboard", "tarjeta madre", "cpu", "ram", "pcie", "sata", "bios", "socket"],
    "cell": ["célula", "celula", "núcleo", "mitocondria", "ribosoma", "retículo", "golgi", "membrana"],
    "heart": ["corazón", "corazon", "ventrículo", "aurícula", "aorta", "válvula"],
    "lungs": ["pulmón", "pulmon", "bronquio", "tráquea", "alvéolo", "respiratorio"],
    "brain": ["cerebro", "neurona", "lóbulo", "cerebelo", "sistema nervioso"],
    "network": ["red", "router", "switch", "firewall", "servidor", "topología", "ethernet"],
}
TOPIC_ES = {
    "motherboard": "placa madre", "cell": "célula", "heart": "corazón", "lungs": "sistema respiratorio",
    "brain": "cerebro", "network": "red informática", "general": "contenido analizado"
}
PARTS = {
    "motherboard": ["Procesador", "Socket del procesador", "Ranuras de memoria RAM", "Chipset", "Conectores SATA", "Ranuras PCIe", "Conector ATX", "Batería CMOS", "Puertos de entrada y salida"],
    "cell": ["Membrana celular", "Citoplasma", "Núcleo", "Nucléolo", "Mitocondrias", "Ribosomas", "Retículo endoplasmático", "Aparato de Golgi"],
    "heart": ["Aurícula derecha", "Aurícula izquierda", "Ventrículo derecho", "Ventrículo izquierdo", "Aorta", "Válvulas", "Arteria pulmonar"],
    "lungs": ["Tráquea", "Bronquios", "Bronquiolos", "Alvéolos", "Pulmón derecho", "Pulmón izquierdo", "Diafragma"],
    "brain": ["Lóbulo frontal", "Lóbulo parietal", "Lóbulo temporal", "Lóbulo occipital", "Cerebelo", "Tronco encefálico"],
    "network": ["Router", "Switch", "Firewall", "Servidor", "Punto de acceso", "Cliente", "Enlace WAN"],
}
CURATED = {
    key: [
        {"title": f"{TOPIC_ES.get(key, key).title()} — modelo educativo", "source": "Sketchfab", "url": f"https://sketchfab.com/search?features=downloadable&type=models&q={key}", "embed_url": "", "thumbnail": "", "license": "Verificar la licencia del autor", "reason": "Búsqueda descargable relacionada con el tema."},
        {"title": f"{TOPIC_ES.get(key, key).title()} — repositorio académico", "source": "Smithsonian/NIH/Web", "url": f"https://www.google.com/search?q={key}+educational+3D+model+GLB", "embed_url": "", "thumbnail": "", "license": "Verificar la licencia del recurso", "reason": "Alternativa institucional o académica para localizar un recurso 3D."},
        {"title": f"{TOPIC_ES.get(key, key).title()} — GLB para web", "source": "Búsqueda GLB", "url": f"https://www.google.com/search?q={key}+GLB+3D+model+download", "embed_url": "", "thumbnail": "", "license": "Verificar la licencia antes de utilizar", "reason": "Alternativa compatible con visores web y realidad aumentada."},
    ] for key in [*TOPICS.keys(), "general"]
}

class ModelCandidate(BaseModel):
    rank: int
    title: str
    source: str
    url: str
    embed_url: str = ""
    thumbnail: str = ""
    license: str
    reason: str
    score: int = 0

class RubricRow(BaseModel):
    criterion: str
    points: int
    excellent: str
    proficient: str
    developing: str
    beginning: str

class DraftResponse(BaseModel):
    project_id: str
    filename: str
    original_file_url: str
    analysis_mode: str
    detected_topic: str
    summary: str
    alt_text: str
    keywords: list[str]
    model_candidates: list[ModelCandidate]
    selected_model_index: int
    activity: dict
    rubric: list[RubricRow]
    quality_report: dict

class BuildRequest(BaseModel):
    project_id: str
    selected_model_index: int = 0
    custom_model_url: str = ""
    custom_model_embed: str = ""
    custom_model_title: str = ""
    activity: dict
    rubric: list[RubricRow]

class BuildResponse(BaseModel):
    scorm_url: str
    validation: dict
    preview_url: str


def normalize_text(value: str) -> str:
    text = re.sub(r"\s+", " ", (value or "")).strip()
    replacements = {
        "SCROM": "SCORM", "Blackbord": "Blackboard", "rubrica": "rúbrica", "evaluacion": "evaluación",
        "identifcar": "identificar", "estudainte": "estudiante", "actvidad": "actividad", "3d": "3D"
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text


def extract_content(path: Path) -> tuple[str, dict]:
    suffix = path.suffix.lower()
    meta: dict = {"type": suffix.lstrip(".") or "archivo"}
    try:
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            meta["pages"] = len(reader.pages)
            text = "\n".join((p.extract_text() or "") for p in reader.pages[:15])
            return text[:30000], meta
        if suffix == ".docx":
            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            meta["paragraphs"] = len(doc.paragraphs)
            return text[:30000], meta
        if suffix in {".txt", ".md", ".csv"}:
            text = path.read_text(encoding="utf-8", errors="ignore")
            return text[:30000], meta
        if suffix in {".html", ".htm"}:
            raw = path.read_text(encoding="utf-8", errors="ignore")
            return BeautifulSoup(raw, "html.parser").get_text(" ")[:30000], meta
        if suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif"}:
            with Image.open(path) as img:
                meta.update({"width": img.width, "height": img.height, "format": img.format, "mode": img.mode})
            return "", meta
    except Exception as exc:
        meta["warning"] = str(exc)
    return "", meta


def detect_topic(text: str, description: str, filename: str) -> str:
    haystack = f"{text} {description} {filename}".lower()
    scores = {topic: sum(haystack.count(term) for term in terms) for topic, terms in TOPICS.items()}
    winner = max(scores, key=scores.get)
    return winner if scores[winner] > 0 else "general"


def extract_keywords(text: str, description: str, topic: str) -> list[str]:
    stop = {"para", "como", "esta", "este", "desde", "sobre", "crear", "actividad", "estudiante", "imagen", "documento", "deberá", "debera", "entre", "donde", "cada", "tiene", "una", "con", "los", "las", "del", "que", "por"}
    words = re.findall(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9-]{4,}", f"{description} {text[:5000]}".lower())
    freq: dict[str, int] = {}
    for word in words:
        if word not in stop:
            freq[word] = freq.get(word, 0) + 1
    keys = [w for w, _ in sorted(freq.items(), key=lambda x: (-x[1], x[0]))[:8]]
    if topic != "general" and TOPIC_ES[topic] not in keys:
        keys.insert(0, TOPIC_ES[topic])
    return keys[:8]


def image_data_url(path: Path) -> str:
    mime = mimetypes.guess_type(path.name)[0] or "image/jpeg"
    return f"data:{mime};base64,{base64.b64encode(path.read_bytes()).decode('ascii')}"


def parse_json_object(raw: str) -> dict:
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?", "", raw).strip()
    raw = re.sub(r"```$", "", raw).strip()
    start, end = raw.find("{"), raw.rfind("}")
    if start < 0 or end < start:
        raise ValueError("La IA no devolvió JSON válido")
    return json.loads(raw[start:end + 1])


def ai_analyze(path: Path, extracted_text: str, description: str, metadata: dict) -> tuple[dict | None, str]:
    if not (AI_BASE_URL and AI_API_KEY and AI_MODEL):
        return None, "local"
    system = (
        "Eres un diseñador instruccional universitario experto en español de Puerto Rico, accesibilidad y Blackboard Ultra. "
        "Analiza cuidadosamente el archivo y la descripción. Corrige ortografía y devuelve exclusivamente JSON válido. "
        "No inventes detalles visuales que no sean visibles o datos que no estén en el documento."
    )
    schema = {
        "detected_topic": "tema breve",
        "summary": "resumen factual de 80 a 140 palabras",
        "alt_text": "texto alternativo preciso; máximo 180 caracteres",
        "keywords": ["6 a 8 palabras clave"],
        "activity": {
            "title": "título",
            "estimated_minutes": 30,
            "bloom_level": "Aplicar o Analizar",
            "objective": "objetivo observable",
            "context": "contexto",
            "instructions": ["4 a 6 pasos"],
            "parts_to_review": ["conceptos o elementos"],
            "question": "pregunta de desarrollo",
            "self_check": [{"question": "pregunta", "options": ["A", "B", "C", "D"], "correct_index": 0, "feedback": "retroalimentación"}]
        },
        "rubric": [{"criterion": "criterio", "points": 25, "excellent": "nivel", "proficient": "nivel", "developing": "nivel", "beginning": "nivel"}]
    }
    prompt = (
        f"Descripción del profesor: {description}\nMetadatos: {json.dumps(metadata, ensure_ascii=False)}\n"
        f"Texto extraído: {extracted_text[:18000]}\n\nCrea una actividad completa y una rúbrica de 100 puntos. "
        f"Incluye exactamente 3 preguntas de autoevaluación. Usa este esquema: {json.dumps(schema, ensure_ascii=False)}"
    )
    content: list[dict] | str = prompt
    if path.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}:
        content = [{"type": "text", "text": prompt}, {"type": "image_url", "image_url": {"url": image_data_url(path)}}]
    payload = {
        "model": AI_MODEL,
        "temperature": 0.2,
        "response_format": {"type": "json_object"},
        "messages": [{"role": "system", "content": system}, {"role": "user", "content": content}],
    }
    try:
        response = requests.post(f"{AI_BASE_URL}/chat/completions", headers={"Authorization": f"Bearer {AI_API_KEY}", "Content-Type": "application/json"}, json=payload, timeout=120)
        response.raise_for_status()
        message = response.json()["choices"][0]["message"]["content"]
        return parse_json_object(message), "external_ai"
    except Exception as exc:
        return {"ai_error": str(exc)}, "fallback_after_ai_error"


def local_activity(topic: str, description: str, text: str, metadata: dict, keywords: list[str]) -> dict:
    topic_es = TOPIC_ES.get(topic, "contenido analizado")
    desc = normalize_text(description)
    summary = normalize_text(text[:900]) if text else (
        f"Recurso visual en formato {metadata.get('format', metadata.get('type', 'imagen'))}, "
        f"con dimensiones de {metadata.get('width', '?')} por {metadata.get('height', '?')} píxeles. "
        "El contenido se interpretará junto con la descripción proporcionada por el profesor."
    )
    parts = PARTS.get(topic, keywords[:6] or ["concepto principal", "evidencia", "aplicación"])
    return {
        "detected_topic": topic_es,
        "summary": summary,
        "alt_text": f"Recurso educativo relacionado con {topic_es}. Consulte las instrucciones para analizar sus elementos principales.",
        "keywords": keywords,
        "activity": {
            "title": f"Análisis aplicado de {topic_es}",
            "estimated_minutes": 35,
            "bloom_level": "Analizar",
            "objective": f"Analizar el recurso relacionado con {topic_es}, identificar sus elementos principales y explicar su función mediante evidencia observable.",
            "context": desc or f"Actividad creada a partir de un recurso sobre {topic_es}.",
            "instructions": [
                "Examine detenidamente la imagen o el documento original.",
                "Revise el modelo 3D seleccionado, cuando esté disponible, y compárelo con el recurso original.",
                "Identifique por lo menos tres elementos o conceptos relevantes.",
                "Explique la función, relación o importancia de cada elemento identificado.",
                "Complete la autoevaluación y redacte su respuesta final en el espacio provisto.",
                "Revise la rúbrica antes de guardar y finalizar la actividad."
            ],
            "parts_to_review": parts,
            "question": f"A partir del recurso presentado, analice tres o más elementos de {topic_es}. Explique la función de cada uno, establezca relaciones entre ellos y sustente su respuesta con evidencia observable en la imagen o el documento.",
            "self_check": [
                {"question": "¿Cuál es el primer paso requerido en esta actividad?", "options": ["Examinar el recurso original", "Ignorar la imagen", "Descargar un programa", "Contestar sin analizar"], "correct_index": 0, "feedback": "Primero se debe examinar cuidadosamente el recurso original."},
                {"question": "¿Cuántos elementos principales debe identificar como mínimo?", "options": ["Uno", "Dos", "Tres", "Diez"], "correct_index": 2, "feedback": "La actividad solicita identificar por lo menos tres elementos."},
                {"question": "¿Qué debe utilizar para sustentar su respuesta?", "options": ["Opiniones sin evidencia", "Evidencia del recurso", "Solo definiciones memorizadas", "Información no relacionada"], "correct_index": 1, "feedback": "La respuesta debe apoyarse en evidencia observable del recurso."}
            ]
        },
        "rubric": default_rubric()
    }


def default_rubric() -> list[dict]:
    return [
        {"criterion": "Identificación de elementos", "points": 20, "excellent": "Identifica cuatro o más elementos con precisión y vocabulario técnico.", "proficient": "Identifica tres elementos correctamente.", "developing": "Identifica dos elementos o presenta imprecisiones.", "beginning": "Identifica uno o ningún elemento correctamente."},
        {"criterion": "Análisis y explicación", "points": 30, "excellent": "Explica funciones y relaciones con profundidad, claridad y precisión.", "proficient": "Explica correctamente las funciones principales y algunas relaciones.", "developing": "Ofrece explicaciones parciales o predominantemente descriptivas.", "beginning": "Presenta explicaciones mínimas, incorrectas o desconectadas."},
        {"criterion": "Uso de evidencia", "points": 20, "excellent": "Integra evidencia específica y pertinente del recurso en toda la respuesta.", "proficient": "Utiliza evidencia pertinente en la mayoría de la respuesta.", "developing": "Incluye referencias generales con poca evidencia específica.", "beginning": "No utiliza evidencia del recurso."},
        {"criterion": "Organización, redacción y ortografía", "points": 15, "excellent": "Respuesta coherente, bien organizada y sin errores significativos.", "proficient": "Respuesta clara con errores menores.", "developing": "La organización o los errores afectan parcialmente la claridad.", "beginning": "La redacción dificulta comprender la respuesta."},
        {"criterion": "Cumplimiento y reflexión", "points": 15, "excellent": "Cumple todos los requisitos e integra una reflexión fundamentada.", "proficient": "Cumple los requisitos principales e incluye una conclusión adecuada.", "developing": "Cumple parcialmente y presenta una conclusión limitada.", "beginning": "No cumple los requisitos esenciales."}
    ]


def normalize_ai_result(result: dict, fallback: dict) -> dict:
    if not result or result.get("ai_error"):
        return fallback
    output = fallback.copy()
    for key in ("detected_topic", "summary", "alt_text", "keywords", "activity", "rubric"):
        if result.get(key):
            output[key] = result[key]
    rubric = output.get("rubric") or default_rubric()
    total = sum(int(r.get("points", 0)) for r in rubric)
    if total != 100:
        rubric = default_rubric()
    output["rubric"] = rubric
    activity = output["activity"]
    activity.setdefault("self_check", fallback["activity"]["self_check"])
    activity["self_check"] = activity["self_check"][:3]
    return output



def normalize_model_candidate(candidate: dict) -> dict:
    """Ensure a Sketchfab model page URL has a usable embed URL."""
    candidate = dict(candidate)
    url = candidate.get("url", "") or ""
    embed = candidate.get("embed_url", "") or ""
    uid = ""
    if "/models/" in embed:
        uid = embed.split("/models/", 1)[1].split("/", 1)[0].split("?", 1)[0]
    if not uid:
        match = re.search(r"-([0-9a-fA-F]{32})(?:[/?#]|$)", url)
        if match:
            uid = match.group(1)
    if not uid:
        match = re.search(r"/models/([0-9a-fA-F]{32})(?:[/?#]|$)", url)
        if match:
            uid = match.group(1)
    if uid:
        candidate["embed_url"] = f"https://sketchfab.com/models/{uid}/embed?autostart=1&ui_theme=dark&ui_infos=0&ui_watermark=0"
        candidate["embeddable"] = True
    else:
        candidate["embeddable"] = bool(embed)
    return candidate


def extract_sketchfab_url_from_embed(value: str) -> str:
    """Extract and validate a Sketchfab model URL from a pasted URL or iframe embed code."""
    raw = (value or "").strip()
    if not raw:
        return ""
    match = re.search(r"src\s*=\s*[\"']([^\"']+)[\"']", raw, flags=re.I)
    candidate = html.unescape(match.group(1) if match else raw).strip()
    parsed = urlparse(candidate)
    host = (parsed.hostname or "").lower()
    if host not in {"sketchfab.com", "www.sketchfab.com"}:
        raise HTTPException(400, "El embed debe pertenecer a Sketchfab.")
    uid_match = re.search(r"/models/([0-9a-fA-F]{32})(?:/embed)?", parsed.path)
    if not uid_match:
        uid_match = re.search(r"-([0-9a-fA-F]{32})(?:[/?#]|$)", candidate)
    if not uid_match:
        raise HTTPException(400, "No se pudo identificar el modelo en el código embed de Sketchfab.")
    uid = uid_match.group(1)
    return f"https://sketchfab.com/models/{uid}/embed?autostart=1&ui_theme=dark&ui_infos=0"


def custom_model_candidate(url: str, title: str = "") -> dict:
    item = normalize_model_candidate({
        "title": normalize_text(title) or "Modelo 3D seleccionado por el profesor",
        "source": "Enlace personalizado",
        "url": url.strip(),
        "embed_url": "",
        "thumbnail": "",
        "license": "Verificar la licencia y los permisos del autor",
        "reason": "Modelo indicado manualmente por el profesor.",
    })
    if not item.get("embeddable"):
        raise HTTPException(status_code=400, detail="El enlace no contiene un identificador válido de modelo Sketchfab. Abra el modelo en Sketchfab y copie la dirección completa de su página.")
    return item

def search_models(topic: str, keywords: list[str]) -> list[dict]:
    results: list[dict] = []
    query = " ".join(keywords[:4]) or TOPIC_ES.get(topic, topic)
    headers = {"Authorization": f"Token {SKETCHFAB_API_TOKEN}"} if SKETCHFAB_API_TOKEN else {}
    try:
        response = requests.get("https://api.sketchfab.com/v3/search", params={"type": "models", "q": query, "downloadable": "true", "count": 6}, headers=headers, timeout=12)
        if response.ok:
            for item in response.json().get("results", []):
                uid = item.get("uid", "")
                images = item.get("thumbnails", {}).get("images", [])
                thumbnail = images[-1].get("url", "") if images else ""
                license_data = item.get("license") or {}
                results.append({
                    "title": normalize_text(item.get("name", "Modelo 3D")),
                    "source": "Sketchfab",
                    "url": item.get("viewerUrl") or f"https://sketchfab.com/3d-models/{uid}",
                    "embed_url": f"https://sketchfab.com/models/{uid}/embed" if uid else "",
                    "thumbnail": thumbnail,
                    "license": license_data.get("label") or "Verificar la licencia del autor",
                    "reason": "Resultado encontrado por coincidencia con el tema y las palabras clave.",
                })
    except Exception:
        pass
    for item in CURATED.get(topic, CURATED["general"]):
        if len(results) >= 3:
            break
        results.append(dict(item))
    unique: list[dict] = []
    seen = set()
    for item in results:
        marker = (item.get("title", ""), item.get("url", ""))
        if marker not in seen:
            seen.add(marker)
            unique.append(normalize_model_candidate(item))
        if len(unique) == 3:
            break
    for idx, item in enumerate(unique, 1):
        item["rank"] = idx
        item["score"] = max(1, 100 - (idx - 1) * 12)
    return unique


def quality_report(result: dict, mode: str, candidates: list[dict], metadata: dict) -> dict:
    rubric_total = sum(int(r.get("points", 0)) for r in result.get("rubric", []))
    activity = result.get("activity", {})
    checks = {
        "recurso_analizado": bool(result.get("summary")),
        "texto_alternativo": bool(result.get("alt_text")),
        "objetivo_observable": bool(activity.get("objective")),
        "instrucciones": len(activity.get("instructions", [])) >= 4,
        "autoevaluacion": len(activity.get("self_check", [])) == 3,
        "rubrica_100_puntos": rubric_total == 100,
        "tres_modelos_3d": len(candidates) == 3,
    }
    return {"score": round(sum(checks.values()) / len(checks) * 100), "checks": checks, "analysis_mode": mode, "metadata": metadata}


def project_path(project_id: str) -> Path:
    return PROJECTS / f"{project_id}.json"


def save_project(data: dict) -> None:
    project_path(data["project_id"]).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def load_project(project_id: str) -> dict:
    path = project_path(project_id)
    if not path.exists():
        raise HTTPException(404, "Proyecto no encontrado")
    return json.loads(path.read_text(encoding="utf-8"))


def resource_block(source: Path, asset_name: str, alt_text: str, extracted_text: str) -> str:
    suffix = source.suffix.lower()
    safe_asset = html.escape(asset_name)
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif"}:
        return f"<img class='original-img' src='assets/{safe_asset}' alt='{html.escape(alt_text)}'>"
    if suffix == ".pdf":
        return f"<iframe class='doc-frame' src='assets/{safe_asset}' title='Documento original'></iframe><p><a href='assets/{safe_asset}' target='_blank'>Abrir PDF en otra pestaña</a></p>"
    excerpt = html.escape(extracted_text[:6000]) if extracted_text else "El documento original está disponible para descarga."
    return f"<div class='document-text'><pre>{excerpt}</pre></div><p><a href='assets/{safe_asset}' target='_blank'>Descargar documento original</a></p>"


def build_scorm(project: dict, selected_index: int, activity: dict, rubric: list[dict]) -> tuple[Path, Path]:
    pid = project["project_id"]
    source = Path(project["stored_path"])
    out = SCORM / pid
    if out.exists():
        shutil.rmtree(out)
    assets = out / "assets"
    assets.mkdir(parents=True)
    asset_name = re.sub(r"[^A-Za-z0-9_.-]+", "_", project["filename"])
    shutil.copy2(source, assets / asset_name)
    candidates = project["model_candidates"]
    selected = candidates[selected_index] if 0 <= selected_index < len(candidates) else candidates[0]
    original = resource_block(source, asset_name, project.get("alt_text", "Recurso educativo original"), project.get("extracted_text", ""))
    instructions = "".join(f"<li>{html.escape(str(x))}</li>" for x in activity.get("instructions", []))
    parts = "".join(f"<li>{html.escape(str(x))}</li>" for x in activity.get("parts_to_review", []))
    rubric_rows = "".join(
        f"<tr><th scope='row'>{html.escape(r['criterion'])}</th><td>{r['points']}</td><td>{html.escape(r['excellent'])}</td><td>{html.escape(r['proficient'])}</td><td>{html.escape(r['developing'])}</td><td>{html.escape(r['beginning'])}</td></tr>" for r in rubric
    )
    checks = activity.get("self_check", [])[:3]
    check_html = ""
    for qi, q in enumerate(checks):
        options = "".join(f"<label class='option'><input type='radio' name='q{qi}' value='{oi}'> {html.escape(str(opt))}</label>" for oi, opt in enumerate(q.get("options", [])))
        check_html += f"<fieldset data-correct='{int(q.get('correct_index',0))}' data-feedback='{html.escape(str(q.get('feedback','')))}'><legend>{qi+1}. {html.escape(str(q.get('question','')))}</legend>{options}<p class='feedback' aria-live='polite'></p></fieldset>"
    if selected.get("embed_url"):
        model_html = f"<iframe class='model-frame' src='{html.escape(selected['embed_url'])}' title='Modelo 3D seleccionado' allow='autoplay; fullscreen; xr-spatial-tracking' allowfullscreen></iframe>"
    else:
        model_html = f"<p>El recurso 3D se abre en una biblioteca externa. Verifique su licencia antes del uso institucional.</p><p><a class='button-link' href='{html.escape(selected.get('url','#'))}' target='_blank' rel='noopener'>Abrir modelo 3D seleccionado</a></p>"
    page = f'''<!doctype html><html lang="es"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>{html.escape(activity.get('title','Actividad UCAN'))}</title>
<style>:root{{--green:#007B5F;--yellow:#FED141;--ink:#17211f;--bg:#f4f7f6}}*{{box-sizing:border-box}}body{{margin:0;font-family:Arial,Helvetica,sans-serif;line-height:1.55;color:var(--ink);background:var(--bg)}}header{{background:#073b35;color:#fff;padding:24px}}main{{max-width:1050px;margin:auto;padding:18px}}.card{{background:#fff;border:1px solid #d9e2df;border-radius:14px;padding:20px;margin:16px 0}}h1,h2{{line-height:1.2}}h2{{color:var(--green)}}.original-img{{display:block;max-width:100%;max-height:680px;margin:auto;border-radius:10px}}.doc-frame,.model-frame{{width:100%;height:560px;border:1px solid #bbb;border-radius:10px}}textarea{{width:100%;min-height:220px;padding:12px;font:inherit}}button,.button-link{{display:inline-block;background:var(--green);color:#fff;border:0;border-radius:8px;padding:11px 15px;font-weight:bold;text-decoration:none;cursor:pointer}}button.secondary{{background:#425b55}}table{{border-collapse:collapse;width:100%;font-size:14px}}th,td{{border:1px solid #bbb;padding:8px;vertical-align:top}}fieldset{{margin:14px 0;border:1px solid #ccd8d4;border-radius:10px}}.option{{display:block;padding:7px}}.feedback{{font-weight:bold}}.status{{padding:12px;border-left:5px solid var(--yellow);background:#fff8d8}}.toolbar{{display:flex;gap:10px;flex-wrap:wrap}}pre{{white-space:pre-wrap}}:focus{{outline:3px solid #1d70b8;outline-offset:2px}}@media print{{button,.toolbar{{display:none}}}}</style>
<script>
let API=null;const KEY='ucan_{pid}';
function findAPI(w){{let tries=0;while(w&&tries<10){{try{{if(w.API)return w.API;if(w.parent===w)break;w=w.parent;}}catch(e){{break}}tries++;}}return null;}}
function getValue(k){{try{{return API?API.LMSGetValue(k):''}}catch(e){{return ''}}}}
function setValue(k,v){{try{{if(API)API.LMSSetValue(k,String(v))}}catch(e){{}}}}
function init(){{API=findAPI(window);if(API){{try{{API.LMSInitialize('')}}catch(e){{}}}}let saved=getValue('cmi.suspend_data')||localStorage.getItem(KEY)||'';if(saved)document.getElementById('answer').value=saved;updateCount();}}
function updateCount(){{const t=document.getElementById('answer').value.trim();document.getElementById('count').textContent=(t?t.split(/\s+/).length:0)+' palabras';}}
function autosave(){{const ans=document.getElementById('answer').value.slice(0,3500);localStorage.setItem(KEY,ans);setValue('cmi.suspend_data',ans);setValue('cmi.comments',ans);if(API)API.LMSCommit('');updateCount();document.getElementById('saveStatus').textContent='Borrador guardado.';}}
function gradeCheck(){{let correct=0;document.querySelectorAll('fieldset').forEach(fs=>{{const picked=fs.querySelector('input:checked');const fb=fs.querySelector('.feedback');if(!picked){{fb.textContent='Seleccione una respuesta.';return}}if(Number(picked.value)===Number(fs.dataset.correct)){{correct++;fb.textContent='Correcto. '+fs.dataset.feedback}}else{{fb.textContent='Revise su respuesta. '+fs.dataset.feedback}}}});document.getElementById('checkScore').textContent='Resultado: '+correct+' de {len(checks)}';}}
function complete(){{autosave();const ans=document.getElementById('answer').value.trim();if(ans.length<80){{document.getElementById('saveStatus').textContent='Amplíe la respuesta antes de completar. Se recomiendan al menos 80 caracteres.';return}}setValue('cmi.core.lesson_status','completed');if(API)API.LMSCommit('');document.getElementById('saveStatus').textContent='Actividad marcada como completada y respuesta guardada.';}}
window.addEventListener('load',init);window.addEventListener('beforeunload',()=>{{autosave();if(API)API.LMSFinish('')}});
</script></head><body><header><h1>{html.escape(activity.get('title','Actividad UCAN'))}</h1><p>UCAN Reality Lab — paquete SCORM 1.2 accesible</p></header><main>
<section class='card'><h2>Objetivo</h2><p>{html.escape(activity.get('objective',''))}</p><p><strong>Nivel cognitivo:</strong> {html.escape(str(activity.get('bloom_level','Analizar')))} · <strong>Tiempo estimado:</strong> {html.escape(str(activity.get('estimated_minutes',30)))} minutos</p></section>
<section class='card'><h2>Recurso original</h2><p>{html.escape(project.get('summary',''))}</p>{original}</section>
<section class='card'><h2>Instrucciones</h2><ol>{instructions}</ol><h3>Elementos a revisar</h3><ul>{parts}</ul></section>
<section class='card'><h2>Modelo 3D seleccionado</h2><p><strong>{html.escape(selected.get('title','Modelo 3D'))}</strong> — {html.escape(selected.get('source',''))}</p><p>{html.escape(selected.get('reason',''))}</p>{model_html}</section>
<section class='card'><h2>Autoevaluación</h2>{check_html}<button type='button' onclick='gradeCheck()'>Comprobar respuestas</button><p id='checkScore' aria-live='polite'></p></section>
<section class='card'><h2>Respuesta del estudiante</h2><p>{html.escape(activity.get('question',''))}</p><label for='answer'><strong>Redacte su respuesta:</strong></label><textarea id='answer' oninput='updateCount()' placeholder='Escriba aquí su análisis...'></textarea><div class='toolbar'><button type='button' class='secondary' onclick='autosave()'>Guardar borrador</button><button type='button' onclick='complete()'>Guardar y completar</button><button type='button' class='secondary' onclick='window.print()'>Imprimir</button></div><p id='count'>0 palabras</p><p id='saveStatus' class='status' aria-live='polite'></p></section>
<section class='card'><h2>Rúbrica</h2><table><thead><tr><th>Criterio</th><th>Puntos</th><th>Excelente</th><th>Competente</th><th>En desarrollo</th><th>Inicial</th></tr></thead><tbody>{rubric_rows}</tbody></table></section>
</main></body></html>'''
    (out / "index.html").write_text(page, encoding="utf-8")
    manifest = f'''<?xml version="1.0" encoding="UTF-8"?>
<manifest identifier="UCAN-{pid}" version="1.2" xmlns="http://www.imsproject.org/xsd/imscp_rootv1p1p2" xmlns:adlcp="http://www.adlnet.org/xsd/adlcp_rootv1p2" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
<metadata><schema>ADL SCORM</schema><schemaversion>1.2</schemaversion></metadata>
<organizations default="ORG1"><organization identifier="ORG1"><title>{html.escape(activity.get('title','Actividad UCAN'))}</title><item identifier="ITEM1" identifierref="RES1"><title>{html.escape(activity.get('title','Actividad UCAN'))}</title></item></organization></organizations>
<resources><resource identifier="RES1" type="webcontent" adlcp:scormtype="sco" href="index.html"><file href="index.html"/><file href="assets/{html.escape(asset_name)}"/></resource></resources>
</manifest>'''
    (out / "imsmanifest.xml").write_text(manifest, encoding="utf-8")
    (out / "activity.json").write_text(json.dumps({"activity": activity, "rubric": rubric, "selected_model": selected}, ensure_ascii=False, indent=2), encoding="utf-8")
    zip_path = SCORM / f"UCAN_SCORM_{pid}.zip"
    if zip_path.exists(): zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in out.rglob("*"):
            if path.is_file(): archive.write(path, path.relative_to(out).as_posix())
    return zip_path, out / "index.html"


def validate_scorm(zip_path: Path) -> dict:
    issues: list[str] = []
    required = {"imsmanifest.xml", "index.html"}
    try:
        with zipfile.ZipFile(zip_path) as archive:
            names = set(archive.namelist())
            missing = required - names
            if missing: issues.append("Faltan archivos: " + ", ".join(sorted(missing)))
            manifest_data = archive.read("imsmanifest.xml")
            root = ET.fromstring(manifest_data)
            if not root.tag.endswith("manifest"): issues.append("El XML no contiene una raíz manifest válida.")
            if not any(n.startswith("assets/") for n in names): issues.append("El recurso original no está incluido en assets/.")
            index = archive.read("index.html").decode("utf-8", errors="ignore")
            for marker in ("LMSInitialize", "LMSSetValue", "LMSCommit", "LMSFinish"):
                if marker not in index: issues.append(f"Falta la llamada SCORM {marker}.")
    except Exception as exc:
        issues.append(str(exc))
    return {"valid": not issues, "issues": issues, "package": zip_path.name, "checked_at": dt.datetime.now().isoformat(timespec="seconds")}

@app.get("/", response_class=HTMLResponse)
def home():
    return (BASE / "static" / "index.html").read_text(encoding="utf-8")

@app.get("/api/health")
def health():
    return {"ok": True, "version": "6.0", "ai_configured": bool(AI_BASE_URL and AI_API_KEY and AI_MODEL), "sketchfab_token": bool(SKETCHFAB_API_TOKEN), "uploads_writable": os.access(UPLOADS, os.W_OK), "scorm_writable": os.access(SCORM, os.W_OK)}

@app.post("/api/analyze", response_model=DraftResponse)
async def analyze(file: UploadFile = File(...), description: str = Form(...)):
    if not file.filename:
        raise HTTPException(400, "Seleccione un archivo.")
    suffix = Path(file.filename).suffix.lower()
    if suffix not in ALLOWED_SUFFIXES:
        raise HTTPException(400, "Formato no permitido.")
    content = await file.read(MAX_UPLOAD + 1)
    if len(content) > MAX_UPLOAD:
        raise HTTPException(413, "El archivo excede 25 MB.")
    pid = uuid.uuid4().hex[:12]
    safe_name = re.sub(r"[^A-Za-z0-9_.-]+", "_", Path(file.filename).name)
    dest = UPLOADS / f"{pid}_{safe_name}"
    dest.write_bytes(content)
    extracted, metadata = extract_content(dest)
    topic = detect_topic(extracted, description, safe_name)
    keywords = extract_keywords(extracted, description, topic)
    fallback = local_activity(topic, description, extracted, metadata, keywords)
    ai_result, mode = ai_analyze(dest, extracted, description, metadata)
    result = normalize_ai_result(ai_result, fallback)
    candidates = search_models(topic, result.get("keywords", keywords))
    report = quality_report(result, mode, candidates, metadata)
    project = {
        "project_id": pid, "filename": safe_name, "stored_path": str(dest), "original_file_url": f"/data/uploads/{dest.name}",
        "description": normalize_text(description), "extracted_text": extracted, "analysis_mode": mode,
        "detected_topic": result["detected_topic"], "summary": result["summary"], "alt_text": result["alt_text"],
        "keywords": result["keywords"], "model_candidates": candidates, "selected_model_index": 0,
        "activity": result["activity"], "rubric": result["rubric"], "quality_report": report,
        "created_at": dt.datetime.now().isoformat(timespec="seconds")
    }
    save_project(project)
    return project

@app.post("/api/build", response_model=BuildResponse)
def build(request: BuildRequest):
    project = load_project(request.project_id)
    if len(request.rubric) == 0 or sum(r.points for r in request.rubric) != 100:
        raise HTTPException(400, "La rúbrica debe sumar exactamente 100 puntos.")
    if not request.activity.get("title") or not request.activity.get("objective") or not request.activity.get("question"):
        raise HTTPException(400, "Complete título, objetivo y pregunta.")
    rubric_data = [r.model_dump() for r in request.rubric]
    custom_source = request.custom_model_embed.strip() or request.custom_model_url.strip()
    if custom_source:
        embed_url = extract_sketchfab_url_from_embed(custom_source)
        custom = custom_model_candidate(embed_url, request.custom_model_title)
        custom["embed_url"] = embed_url
        custom["url"] = embed_url.split("/embed", 1)[0]
        custom["source"] = "Sketchfab — embed proporcionado por el profesor"
        custom["reason"] = "Modelo seleccionado y confirmado por el profesor mediante el código embed de Sketchfab."
        project["model_candidates"] = [custom] + project.get("model_candidates", [])
        selected_index = 0
    else:
        selected_index = request.selected_model_index
    zip_path, preview = build_scorm(project, selected_index, request.activity, rubric_data)
    validation = validate_scorm(zip_path)
    project.update({"activity": request.activity, "rubric": rubric_data, "selected_model_index": selected_index, "scorm_url": f"/data/scorm/{zip_path.name}", "validation": validation})
    save_project(project)
    return {"scorm_url": f"/data/scorm/{zip_path.name}", "validation": validation, "preview_url": f"/data/scorm/{project['project_id']}/index.html"}

@app.get("/api/project/{project_id}")
def get_project(project_id: str):
    return load_project(project_id)

@app.get("/api/scorm/{name}")
def download_scorm(name: str):
    safe = Path(name).name
    path = SCORM / safe
    if not path.exists(): raise HTTPException(404, "Paquete no encontrado")
    return FileResponse(path, filename=safe, media_type="application/zip")
