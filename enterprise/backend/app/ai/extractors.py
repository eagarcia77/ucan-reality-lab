from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile
import io
import re

from bs4 import BeautifulSoup
from docx import Document
from PIL import Image
from pypdf import PdfReader


@dataclass(slots=True)
class ExtractedContent:
    text: str
    media_type: str
    metadata: dict
    warnings: list[str]


def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def extract_file(path: Path) -> ExtractedContent:
    suffix = path.suffix.lower()
    warnings: list[str] = []

    if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
        with Image.open(path) as image:
            return ExtractedContent(
                text="",
                media_type=Image.MIME.get(image.format, "image/jpeg"),
                metadata={
                    "width": image.width,
                    "height": image.height,
                    "format": image.format,
                    "mode": image.mode,
                },
                warnings=warnings,
            )

    if suffix == ".pdf":
        reader = PdfReader(str(path))
        pages = [(page.extract_text() or "") for page in reader.pages[:30]]
        if not any(page.strip() for page in pages):
            warnings.append("El PDF no contiene texto extraíble; puede requerir OCR o análisis visual.")
        return ExtractedContent(
            text=_clean("\n".join(pages))[:60000],
            media_type="application/pdf",
            metadata={"pages": len(reader.pages)},
            warnings=warnings,
        )

    if suffix == ".docx":
        document = Document(str(path))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        tables = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        return ExtractedContent(
            text=_clean("\n".join(paragraphs + tables))[:60000],
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            metadata={"paragraphs": len(document.paragraphs), "tables": len(document.tables)},
            warnings=warnings,
        )

    if suffix == ".pptx":
        texts: list[str] = []
        slide_count = 0
        with ZipFile(path) as archive:
            slide_names = sorted(
                name for name in archive.namelist()
                if name.startswith("ppt/slides/slide") and name.endswith(".xml")
            )
            slide_count = len(slide_names)
            for name in slide_names:
                xml = archive.read(name).decode("utf-8", errors="ignore")
                texts.extend(re.findall(r"<a:t>(.*?)</a:t>", xml))
        return ExtractedContent(
            text=_clean(" ".join(BeautifulSoup(t, "html.parser").get_text(" ") for t in texts))[:60000],
            media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            metadata={"slides": slide_count},
            warnings=warnings,
        )

    if suffix in {".txt", ".md", ".csv", ".html", ".htm"}:
        raw = path.read_text(encoding="utf-8", errors="ignore")
        if suffix in {".html", ".htm"}:
            raw = BeautifulSoup(raw, "html.parser").get_text(" ")
        return ExtractedContent(
            text=_clean(raw)[:60000],
            media_type="text/plain",
            metadata={"characters": len(raw)},
            warnings=warnings,
        )

    raise ValueError(f"Tipo de archivo no compatible: {suffix or 'sin extensión'}")
